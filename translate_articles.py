"""Translate PDFs in ``to_translate/`` into Hebrew DOCX files.

For each PDF the pipeline:

1. Computes a SHA-256 hash and looks it up in ``translation_log.json`` so
   the same content is never translated twice (even if filenames differ).
2. Extracts page text and detects whether the PDF contains embedded images.
3. Detects the source language. PDFs that already look Hebrew are skipped
   (marked ``already_hebrew``) and only their original is mirrored.
4. Translates the text page-by-page in <=4500 char chunks via free Google
   Translate (``deep_translator``).
5. Writes a Hebrew DOCX (right-to-left) alongside a copy of the original
   PDF into ``translated/<basename>/``.
6. On failure, moves the original PDF into ``error_translate/`` and records
   the error in the log.
7. Regenerates ``index.html`` with an Original | Translated | Has pictures
   table.

Usage:
    py -3.13 translate_articles.py            # process everything pending
    py -3.13 translate_articles.py --limit 1  # just one file (for testing)
    py -3.13 translate_articles.py --only "Gopher et al. 2001*"  # glob match
"""

from __future__ import annotations

import argparse
import fnmatch
import hashlib
import html
import json
import os
import shutil
import sys
import time
import traceback
import warnings
from dataclasses import dataclass, field
from pathlib import Path

import urllib3
import requests

warnings.simplefilter("ignore", urllib3.exceptions.InsecureRequestWarning)

# --- TLS bypass for the corporate proxy (same trick as _download_drive.py) ---
_orig_session_request = requests.Session.request


def _no_verify_session_request(self, method, url, **kwargs):
    kwargs["verify"] = False
    return _orig_session_request(self, method, url, **kwargs)


requests.Session.request = _no_verify_session_request

_orig_top_request = requests.api.request


def _no_verify_top_request(method, url, **kwargs):
    kwargs["verify"] = False
    return _orig_top_request(method, url, **kwargs)


requests.api.request = _no_verify_top_request


import fitz  # PyMuPDF  # noqa: E402
from deep_translator import GoogleTranslator  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.shared import Pt  # noqa: E402
from langdetect import DetectorFactory, detect  # noqa: E402

DetectorFactory.seed = 0  # deterministic language detection


ROOT = Path(__file__).resolve().parent
SRC_DIR = ROOT / "to_translate"
OUT_DIR = ROOT / "translated"
ERR_DIR = ROOT / "error_translate"
LOG_PATH = ROOT / "translation_log.json"
INDEX_PATH = ROOT / "index.html"

# Google Translate caps a single request at ~5000 characters. Stay under that
# with some headroom for URL-encoding overhead.
CHUNK_CHAR_LIMIT = 4500


# ---------------------------------------------------------------------------
# Log helpers
# ---------------------------------------------------------------------------


@dataclass
class LogEntry:
    sha256: str
    original_name: str
    status: str  # "translated" | "already_hebrew" | "error"
    has_pictures: bool = False
    page_count: int = 0
    source_lang: str | None = None
    translated_relpath: str | None = None  # relative to repo root
    original_relpath: str | None = None  # relative to repo root
    error: str | None = None
    duplicates: list[str] = field(default_factory=list)
    translated_at: str | None = None

    def to_dict(self) -> dict:
        return {
            "sha256": self.sha256,
            "original_name": self.original_name,
            "status": self.status,
            "has_pictures": self.has_pictures,
            "page_count": self.page_count,
            "source_lang": self.source_lang,
            "translated_relpath": self.translated_relpath,
            "original_relpath": self.original_relpath,
            "error": self.error,
            "duplicates": self.duplicates,
            "translated_at": self.translated_at,
        }

    @classmethod
    def from_dict(cls, data: dict) -> "LogEntry":
        return cls(
            sha256=data["sha256"],
            original_name=data["original_name"],
            status=data["status"],
            has_pictures=data.get("has_pictures", False),
            page_count=data.get("page_count", 0),
            source_lang=data.get("source_lang"),
            translated_relpath=data.get("translated_relpath"),
            original_relpath=data.get("original_relpath"),
            error=data.get("error"),
            duplicates=list(data.get("duplicates", [])),
            translated_at=data.get("translated_at"),
        )


def load_log() -> dict[str, LogEntry]:
    if not LOG_PATH.exists():
        return {}
    with LOG_PATH.open("r", encoding="utf-8") as fh:
        raw = json.load(fh)
    return {sha: LogEntry.from_dict(entry) for sha, entry in raw.items()}


def save_log(log: dict[str, LogEntry]) -> None:
    serializable = {sha: entry.to_dict() for sha, entry in log.items()}
    with LOG_PATH.open("w", encoding="utf-8") as fh:
        json.dump(serializable, fh, ensure_ascii=False, indent=2, sort_keys=True)


# ---------------------------------------------------------------------------
# PDF + translation primitives
# ---------------------------------------------------------------------------


def sha256_of(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for block in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def extract_pages_and_images(pdf_path: Path) -> tuple[list[str], bool, int]:
    """Return (per-page text, has_images, page_count)."""
    pages: list[str] = []
    has_images = False
    with fitz.open(pdf_path) as doc:
        for page in doc:
            pages.append(page.get_text("text") or "")
            if not has_images and page.get_images(full=False):
                has_images = True
        page_count = doc.page_count
    return pages, has_images, page_count


def detect_language(sample_text: str) -> str | None:
    sample = (sample_text or "").strip()
    if len(sample) < 40:
        return None
    try:
        return detect(sample[:4000])
    except Exception:
        return None


def chunk_text(text: str, limit: int = CHUNK_CHAR_LIMIT) -> list[str]:
    """Split text into <=limit char chunks on paragraph/sentence boundaries."""
    if len(text) <= limit:
        return [text] if text else []

    chunks: list[str] = []
    paragraphs = text.split("\n")
    buf = ""
    for para in paragraphs:
        candidate = (buf + "\n" + para) if buf else para
        if len(candidate) <= limit:
            buf = candidate
            continue
        if buf:
            chunks.append(buf)
            buf = ""
        # paragraph itself may be too long; sentence-split it
        if len(para) <= limit:
            buf = para
            continue
        sentences = para.replace("? ", "?|").replace("! ", "!|").replace(". ", ".|").split("|")
        for sent in sentences:
            cand = (buf + " " + sent) if buf else sent
            if len(cand) <= limit:
                buf = cand
            else:
                if buf:
                    chunks.append(buf)
                # if a single "sentence" is still too long, hard-wrap it
                while len(sent) > limit:
                    chunks.append(sent[:limit])
                    sent = sent[limit:]
                buf = sent
    if buf:
        chunks.append(buf)
    return chunks


class _Translator:
    """Thin wrapper that retries Google Translate on transient failures."""

    def __init__(self, target: str = "iw"):
        # deep-translator uses "iw" (legacy) or "he" depending on version.
        # GoogleTranslator accepts both via its language map.
        self.target = target
        self._engine = GoogleTranslator(source="auto", target=target)

    def translate(self, text: str, retries: int = 3, backoff: float = 2.0) -> str:
        if not text or not text.strip():
            return ""
        last_exc: Exception | None = None
        for attempt in range(retries):
            try:
                return self._engine.translate(text) or ""
            except Exception as exc:  # noqa: BLE001
                last_exc = exc
                wait = backoff * (attempt + 1)
                print(f"    translate retry {attempt + 1}/{retries} after {wait:.1f}s ({exc})", flush=True)
                time.sleep(wait)
        raise RuntimeError(f"translation failed after {retries} attempts: {last_exc}")


def translate_pages(pages: list[str], translator: _Translator) -> list[str]:
    translated_pages: list[str] = []
    for page_idx, page_text in enumerate(pages, start=1):
        page_text = (page_text or "").strip()
        if not page_text:
            translated_pages.append("")
            continue
        chunks = chunk_text(page_text)
        out_parts: list[str] = []
        for chunk_idx, chunk in enumerate(chunks, start=1):
            print(
                f"    page {page_idx}/{len(pages)} chunk {chunk_idx}/{len(chunks)} ({len(chunk)} chars)",
                flush=True,
            )
            out_parts.append(translator.translate(chunk))
        translated_pages.append("\n".join(out_parts))
    return translated_pages


# ---------------------------------------------------------------------------
# DOCX writer (RTL Hebrew)
# ---------------------------------------------------------------------------


def _set_rtl(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _set_run_rtl(run) -> None:
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)


def write_hebrew_docx(
    out_path: Path,
    title: str,
    pages: list[str],
    metadata: dict,
) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Make the default paragraph format RTL.
    title_p = doc.add_heading(title, level=1)
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    _set_rtl(title_p)
    for run in title_p.runs:
        _set_run_rtl(run)

    meta_lines = [
        f"מסמך מקור: {metadata.get('original_name', '')}",
        f"שפת מקור: {metadata.get('source_lang') or 'לא זוהתה'}",
        f"מספר עמודים: {metadata.get('page_count', 0)}",
        f"מכיל תמונות: {'כן' if metadata.get('has_pictures') else 'לא'}",
        f"תורגם ב-: {metadata.get('translated_at', '')}",
    ]
    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        _set_rtl(p)
        for run in p.runs:
            _set_run_rtl(run)

    doc.add_paragraph("")

    for i, page_text in enumerate(pages, start=1):
        header = doc.add_heading(f"עמוד {i}", level=2)
        header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        _set_rtl(header)
        for run in header.runs:
            _set_run_rtl(run)

        if not page_text.strip():
            empty = doc.add_paragraph("(אין טקסט בעמוד זה — ייתכן ומדובר בעמוד תמונה)")
            empty.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            _set_rtl(empty)
            for run in empty.runs:
                _set_run_rtl(run)
            continue

        for paragraph_text in page_text.split("\n"):
            paragraph_text = paragraph_text.strip()
            if not paragraph_text:
                continue
            p = doc.add_paragraph(paragraph_text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            _set_rtl(p)
            for run in p.runs:
                _set_run_rtl(run)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


# ---------------------------------------------------------------------------
# HTML index
# ---------------------------------------------------------------------------


def regenerate_index(log: dict[str, LogEntry]) -> None:
    rows = []
    entries = sorted(log.values(), key=lambda e: e.original_name.lower())
    for entry in entries:
        original_link = (
            f'<a href="{html.escape(entry.original_relpath)}">{html.escape(entry.original_name)}</a>'
            if entry.original_relpath
            else html.escape(entry.original_name)
        )
        if entry.status == "translated" and entry.translated_relpath:
            translated_link = (
                f'<a href="{html.escape(entry.translated_relpath)}">DOCX</a>'
            )
        elif entry.status == "already_hebrew":
            translated_link = "<em>כבר בעברית</em>"
        elif entry.status == "error":
            err = html.escape(entry.error or "unknown error")
            translated_link = f'<span class="err">שגיאה: {err}</span>'
        else:
            translated_link = "—"
        has_pics = "כן" if entry.has_pictures else "לא"
        rows.append(
            f"<tr><td>{original_link}</td><td>{translated_link}</td>"
            f"<td>{has_pics}</td><td>{entry.page_count}</td>"
            f"<td>{html.escape(entry.source_lang or '')}</td>"
            f"<td>{html.escape(entry.translated_at or '')}</td></tr>"
        )

    body = "\n".join(rows) if rows else (
        "<tr><td colspan='6'>אין רשומות עדיין</td></tr>"
    )

    page = f"""<!DOCTYPE html>
<html lang="he" dir="rtl">
<head>
<meta charset="utf-8">
<title>מאמרים מתורגמים</title>
<style>
  body {{ font-family: Arial, sans-serif; margin: 24px; }}
  h1 {{ font-size: 22px; }}
  table {{ border-collapse: collapse; width: 100%; }}
  th, td {{ border: 1px solid #ccc; padding: 8px; text-align: right; vertical-align: top; }}
  th {{ background: #f4f4f4; }}
  tr:nth-child(even) td {{ background: #fafafa; }}
  a {{ color: #1a4f8b; text-decoration: none; }}
  a:hover {{ text-decoration: underline; }}
  .err {{ color: #b00020; }}
</style>
</head>
<body>
<h1>מאמרים מתורגמים</h1>
<p>סה"כ רשומות: {len(entries)}</p>
<table>
<thead>
<tr>
  <th>מסמך מקור</th>
  <th>תרגום</th>
  <th>מכיל תמונות</th>
  <th>מספר עמודים</th>
  <th>שפת מקור</th>
  <th>תורגם ב-</th>
</tr>
</thead>
<tbody>
{body}
</tbody>
</table>
</body>
</html>
"""
    INDEX_PATH.write_text(page, encoding="utf-8")


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------


def safe_folder_name(name: str) -> str:
    stem = Path(name).stem
    out = []
    for ch in stem:
        if ch in '<>:"/\\|?*':
            out.append("_")
        else:
            out.append(ch)
    cleaned = "".join(out).strip().strip(".")
    return cleaned or "document"


def process_pdf(
    pdf_path: Path,
    log: dict[str, LogEntry],
    translator: _Translator,
) -> LogEntry:
    print(f"\n=== {pdf_path.name} ===", flush=True)
    digest = sha256_of(pdf_path)

    existing = log.get(digest)
    if existing is not None:
        if existing.original_name != pdf_path.name and pdf_path.name not in existing.duplicates:
            existing.duplicates.append(pdf_path.name)
        print(f"  already in log as '{existing.original_name}' (status={existing.status}); skipping.", flush=True)
        return existing

    pages, has_pictures, page_count = extract_pages_and_images(pdf_path)
    full_text = "\n".join(pages)
    src_lang = detect_language(full_text)
    print(f"  pages={page_count} has_pictures={has_pictures} detected_lang={src_lang}", flush=True)

    if src_lang in {"he", "iw"}:
        dest_folder = OUT_DIR / safe_folder_name(pdf_path.name)
        dest_folder.mkdir(parents=True, exist_ok=True)
        dest_pdf = dest_folder / pdf_path.name
        shutil.copy2(pdf_path, dest_pdf)
        entry = LogEntry(
            sha256=digest,
            original_name=pdf_path.name,
            status="already_hebrew",
            has_pictures=has_pictures,
            page_count=page_count,
            source_lang=src_lang,
            translated_relpath=None,
            original_relpath=str(dest_pdf.relative_to(ROOT)).replace("\\", "/"),
            translated_at=time.strftime("%Y-%m-%d %H:%M:%S"),
        )
        log[digest] = entry
        print("  marked as already_hebrew; original copied to translated/.", flush=True)
        return entry

    try:
        translated_pages = translate_pages(pages, translator)
    except Exception as exc:  # noqa: BLE001
        ERR_DIR.mkdir(parents=True, exist_ok=True)
        err_dest = ERR_DIR / pdf_path.name
        try:
            shutil.move(str(pdf_path), str(err_dest))
            err_rel = str(err_dest.relative_to(ROOT)).replace("\\", "/")
        except Exception:
            err_rel = None
        traceback.print_exc()
        entry = LogEntry(
            sha256=digest,
            original_name=pdf_path.name,
            status="error",
            has_pictures=has_pictures,
            page_count=page_count,
            source_lang=src_lang,
            translated_relpath=None,
            original_relpath=err_rel,
            error=str(exc),
            translated_at=time.strftime("%Y-%m-%d %H:%M:%S"),
        )
        log[digest] = entry
        print(f"  FAILED: {exc}; moved to error_translate/.", flush=True)
        return entry

    dest_folder = OUT_DIR / safe_folder_name(pdf_path.name)
    dest_folder.mkdir(parents=True, exist_ok=True)
    dest_pdf = dest_folder / pdf_path.name
    shutil.copy2(pdf_path, dest_pdf)

    docx_name = Path(pdf_path.name).stem + ".he.docx"
    dest_docx = dest_folder / docx_name

    metadata = {
        "original_name": pdf_path.name,
        "source_lang": src_lang,
        "page_count": page_count,
        "has_pictures": has_pictures,
        "translated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    write_hebrew_docx(dest_docx, Path(pdf_path.name).stem, translated_pages, metadata)

    entry = LogEntry(
        sha256=digest,
        original_name=pdf_path.name,
        status="translated",
        has_pictures=has_pictures,
        page_count=page_count,
        source_lang=src_lang,
        translated_relpath=str(dest_docx.relative_to(ROOT)).replace("\\", "/"),
        original_relpath=str(dest_pdf.relative_to(ROOT)).replace("\\", "/"),
        translated_at=metadata["translated_at"],
    )
    log[digest] = entry
    print(f"  wrote {dest_docx.relative_to(ROOT)}", flush=True)
    return entry


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--limit", type=int, default=None, help="Process at most N pending files")
    parser.add_argument("--only", type=str, default=None, help="Glob pattern; only matching filenames are processed")
    parser.add_argument("--reset-errors", action="store_true", help="Retry files previously marked as error")
    args = parser.parse_args()

    if not SRC_DIR.exists():
        print(f"Source dir not found: {SRC_DIR}", file=sys.stderr)
        return 2

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ERR_DIR.mkdir(parents=True, exist_ok=True)

    log = load_log()

    if args.reset_errors:
        for sha, entry in list(log.items()):
            if entry.status == "error":
                del log[sha]

    translator = _Translator(target="iw")

    pdfs = sorted(p for p in SRC_DIR.iterdir() if p.is_file() and p.suffix.lower() == ".pdf")
    if args.only:
        pdfs = [p for p in pdfs if fnmatch.fnmatch(p.name, args.only)]

    processed = 0
    try:
        for pdf in pdfs:
            digest = sha256_of(pdf)
            if digest in log and log[digest].status != "error":
                # Already handled; still surface in log + index but don't re-translate.
                process_pdf(pdf, log, translator)  # adds to duplicates / no-ops
                continue
            process_pdf(pdf, log, translator)
            processed += 1
            if args.limit is not None and processed >= args.limit:
                print(f"\nHit --limit={args.limit}; stopping.", flush=True)
                break
    finally:
        save_log(log)
        regenerate_index(log)
        print(f"\nLog: {LOG_PATH.relative_to(ROOT)}", flush=True)
        print(f"Index: {INDEX_PATH.relative_to(ROOT)}", flush=True)

    return 0


if __name__ == "__main__":
    sys.exit(main())
